import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_gdd():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Colors
    DARK_RED = RGBColor(0x8B, 0x00, 0x00)
    GOLD = RGBColor(0xB8, 0x86, 0x0B)
    DARK_BLUE = RGBColor(0x1B, 0x36, 0x5D)
    GRAY = RGBColor(0x55, 0x55, 0x55)

    # --- Title Header ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("KAGE NO YOKAI: 1200 AD\n(Тень Ёкаев: 1200 год)")
    run_title.font.name = 'Georgia'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_RED

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(20)
    run_sub = p_sub.add_run("Полный Дизайн-Документ (Game Design Document / GDD)\n3D-Pixel Action-Roguelike в сеттинге Средневековой Японии — 12 Кампаний и Полный Разбор")
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = GRAY

    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Georgia'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = DARK_RED
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Georgia'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = DARK_BLUE
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = GOLD
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        p.add_run(text)
        return p

    # Section 1
    add_h1("1. Обзор Концепции и Ключевые Механики")
    p = doc.add_paragraph()
    p.add_run("Kage no Yokai: 1200 AD — динамичный экшен-рогалик (Action-Roguelike / Horde Survival) с видом от 3-го лица в гибридном стиле 3D-Pixel Art. Игрок выбирает одного из японских воинов эпохи Камакура и сражается с армией японских демонов (ёкаев) через 12 сюжетных локаций с уникальными заставками, волновым спавном и епическими боссами.")

    add_h2("Ключевые фичи:")
    add_bullet("Управление: ", "WASD для перемещения персонажа, Мышь для направления прицела и атаки/стрельбы.")
    add_bullet("Графический Стиль: ", "50/50 3D Pixel Art (low-poly 3D объекты + пиксельные текстуры/билбординг + атмосфера древней Японии).")
    add_bullet("12 Игровых Уровней: ", "Каждый уровень имеет уникальную заставку при загрузке, свое визуальное оформление, уникальный набор врагов и легендарного Босса.")
    add_bullet("Мета-прогрессия: ", "Получение XP за уничтожение ёкаев в конце забега. Повышение уровня аккаунта разблокирует новые скины персонажей.")
    add_bullet("Система Предметов: ", "До 3 слотов оружия и 3 слотов пассивных способностей. 6 степеней редкости предметов (от Обычного до Божественного). Реролл и Скип при левел-апе.")

    # Section 2: 12 Detailed Levels
    add_h1("2. Полный Детализированный Список 12 Уровней (Этапов)")
    p = doc.add_paragraph()
    p.add_run("Каждый из 12 уровней открывается поочередно. Перед началом уровня выводится живописная заставка-экран с названием главы, легендой и атмосферным артом.")

    levels = [
        ("Этап 1: Разрушенные Врата Камакуры (Kamakura Gate)",
         "«Пепел и Кровь». Горящие японские пагоды, разрушенные глиняные стены, падающий пепел сакуры.",
         "Гаки (голодные бесы), Тёчин-обакэ (фонари-огнееды), Дикие Каппы.",
         "Они-Военачальник (Aka-Oni) — Огромный демон с шипованной палицей Канабо, вызывающий трещины в земле."),

        ("Этап 2: Бамбуковый Лес Затмения (Eclipse Bamboo Grove)",
         "«Шорох во Тьме». Густые лунные бамбуковые рощи, густой синий туман, мерцание светлячков-хитодама.",
         "Карасу-тэнгу (воины-вороны), Нукэкуби (летающие головы), Призрачные волки.",
         "Владыка Тэнгу (Karasu-Tengu Chief) — Развивает ураганные крылья, пикирует с небес и мечет перья-клинки."),

        ("Этап 3: Оскверненная Деревня Киото (Defiled Kyoto Village)", "«Заброшенные Очаги». Узкие деревенские улочки, мосты над алой рекой, заброшенные дома с соломенными крышами.", "Дзёрогумо (паучьи девы), Баконэко (двухвостые демонические коты), Злые Каппы.", "Мать Пауков Дзёрогумо (Spider Empress) — Заплетает арену паутиной и призывает сотни ядовитых паучат."),

        ("Этап 4: Горный Храм Снега (Snowy Mountain Shrine)", "«Ледяная Тишина». Высокогорный буддийский монастырь, заснеженные тории, морозный ветер и замерзшие водопады.", "Юки-онна (снежные девы), Снежные Гаки, Морозные бесы.", "Ледяная Императрица Юки-онна (Frost Queen) — Замораживает землю под ногами и вызывает снежную бурю."),

        ("Этап 5: Кладбище Древних Самураев (Samurai Graveyard)", "«Шепот Проклятых». Каменные надгробия, разрушенные торо (фонари), мистический зеленый туман и призраки.", "Духи тёмных самураев, Проклятые доспехи, Летающие черепа.", "Гашадокуро (Gashadokuro) — Исполинский skeleton высотой 15 метров, пытающийся раздавить героя костяной лапой."),

        ("Этап 6: Топь Тёмных Духов (Swamp of Damned Souls)", "«Гнилые Воды». Заболоченное озеро с ядовитыми лотосами, гниющие ворота тории и пузырящийся газ.", "Мастера Каппа, Ядовитые слизни, Водяные кицунэ.", "Владыка Глубин Мэдзути (Mizuchi Naga) — Огромный змееподобный водяной дракон, плюющийся ядовитой кислотой."),

        ("Этап 7: Затопленный Замок Синоби (Sunken Shinobi Castle)", "«Тени на Крышах». Многоярусный замковый комплекс под покровом ночи, черепичные крыши и секретные ходы.", "Теневые Ниндзя-ёкаи, Летающие сюрикены-духи, Нукэкуби.", "Глава Клана Ночных Теней (Shadow Shinobi Master) — Телепортируется, создает 5 иллюзорных клонов."),

        ("Этап 8: Вулканическое Ущелье Кагуцути (Kagutsuchi Volcano)", "«Огненная Бездна». Скалистые тропы вокруг кипящих лавовых рек, падающие метеориты и раскаленный пепел.", "Огненные бесы, Каша (демоны-кошки на горящих колесах), Лавовые големы.", "Каша — Демоница Пламени (Kasha Wheel Demon) — Мчится по арене на горящей колеснице, оставляя пламенный след."),

        ("Этап 9: Пагода Тысячи Молитв (Pagoda of 1000 Prayers)", "«Священное Осквернение». Внутренние залы древней пагоды, горящие свечи, парящие священные свитки.", "Оскверненные Монахи, Духи свитков, Золотые демоны.", "Великий Дайтэнгу (Daitengu High Priest) — Колдует священные заклинания огня и воздуха в замкнутом зале."),

        ("Этап 10: Озеро Багряной Луны (Lake of the Crimson Moon)", "«Кровавый Глянец». Гладкое зеркальное озеро, отражающее огромную красную луну и лепестки сакуры.", "Девятихвостые Кицунэ, Иллюзорные лисицы, Призрачные самураи.", "Девятихвостая Кицунэ (Kyubi no Kitsune) — Призывает огненные сферы, ослепляет вспышками и оборачивается туманом."),

        ("Этап 11: Врата Расёмон — Порог Преисподней (Gates of Rashomon)", "«Раскол Миров». Гигантские древние врата Расёмон, разрываемые фиолетово-алым порталом в Ёми.", "Элитная гвардия демонов, Двойные бесы, Тёмные драконы.", "Близнецы Райдзин и Фудзин (Thunder & Wind Gods) — Двойной босс: Повелитель Молнии и Повелитель Урагана."),

        ("Этап 12: Сердце Царства Ёми (Heart of Yomi Realm)", "«Абсолютная Тьма». Финальная арена в ином измерении с алыми кристаллами и парящими обломками скал.", "Все виды высших ёкаев и теневые копии боссов.", "ФИНАЛЬНЫЙ БОСС: Ямата-но Ороти (Yamata-no-Orochi) — Восьмиглавый древний дракон хаоса с 8 уникальными атаками.")
    ]

    for title, desc, enemies, boss in levels:
        add_h2(title)
        add_bullet("Заставка и Атмосфера: ", desc)
        add_bullet("Враги на уровне: ", enemies)
        add_bullet("БОСС ЭТАПА: ", boss)

    # Section 3: Skins
    add_h1("3. Персонажи, Скины и Прогрессия")
    p = doc.add_paragraph()
    p.add_run("В меню игрок настраивает персонажа. Первые 3 скина доступны сразу, последующие открываются за XP профиля:")
    add_bullet("1. Самурай (Катана): ", "Базовый скин. Высокий шанс критического удара.")
    add_bullet("2. Ниндзя / Синоби (Сюрикены): ", "Базовый скин. Скорость передвижения, рывок с бессмертием.")
    add_bullet("3. Лучник (Лук Юми): ", "Базовый скин. Дальнобойный урон, пробитие 1 цели.")
    add_bullet("4. Самурай Нодати (Нодати): ", "Уровень 5. Огромный меч с широким замахом и отбрасыванием.")
    add_bullet("5. Жрица Мико (Печати Офуда): ", "Уровень 10. Святой урон по площади, печать исцеления.")
    add_bullet("6. Монах Сохэй (Нагината): ", "Уровень 15. Круговая защита и регулярный блок урона.")

    # Section 4: Upgrades & Rarities
    add_h1("4. Карточки Прокачки, 6 Степеней Редкости и Синергии")
    p = doc.add_paragraph()
    p.add_run("При левел-апе дается 3 карточки на выбор, 1 реролл и 1 кнопка пропуска. Ограничение: максимум 3 оружия и 3 способности.")
    
    add_h2("6 Степеней Редкости:")
    add_bullet("⚪ Обычное (Common - 60%): ", "Базовые параметры.")
    add_bullet("🟢 Необычное (Uncommon - 25%): ", "+25% характеристик, +1 снаряд.")
    add_bullet("🔵 Редкое (Rare - 10%): ", "+50% характеристик, эффекты кровотечения.")
    add_bullet("🟣 Эпическое (Epic - 4%): ", "+100% характеристик, +2 отскока/взмаха.")
    add_bullet("🟡 Легендарное (Legendary - 0.9%): ", "+200% урон, уникальные элементы.")
    add_bullet("🔴 Божественное (Divine - 0.1%): ", "+400% урон, супер-комбо.")

    add_h2("Примеры Синергий (Комбо-Сеты):")
    add_bullet("Пламенный Ураган (Inferno Vortex): ", "Огненное Дыхание + Вихрь Клинков.")
    add_bullet("Токсичный Шок (Toxic Shock): ", "Ядовитый Туман + Громовой Удар.")
    add_bullet("Призрак Древних (Spectral Host): ", "Печать Душ + Ледяная Аура + Дух Волка (Призыв самураев-духов).")
    add_bullet("Стрела Райдзина (Raijin's Bolt): ", "Лук Юми + Громовой Удар (Удары молний от стрел).")

    # Section 5: Technical Plan
    add_h1("5. План Разработки Прототипа на HTML5 (Three.js)")
    p = doc.add_paragraph()
    p.add_run("Выбран стек HTML5 + Three.js + Vite. ИИ (я) берет на себя полностью генерацию процедурных 3D-моделей (самураи, ёкаи, бамбук, храмы), создание 3D pixel shaders, синтез звуковых эффектов (Web Audio API), интерфейса и всей игровой механики.")

    doc.save("c:/Users/dorel/OneDrive/Desktop/RogueLike_game/GDD_Yokai_Roguelike.docx")
    print("Updated 12-level GDD Docx saved!")

if __name__ == "__main__":
    create_gdd()

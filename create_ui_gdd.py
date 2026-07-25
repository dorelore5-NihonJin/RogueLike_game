import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_name}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

DARK_RED = RGBColor(0x8B, 0x00, 0x00)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK_BLUE = RGBColor(0x1B, 0x36, 0x5D)
GRAY = RGBColor(0x55, 0x55, 0x55)

def create_base_doc(title_text, subtitle_text):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title_text)
    r_title.font.name = 'Georgia'
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = DARK_RED

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run(subtitle_text)
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = GRAY
    return doc

def add_h1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = 'Georgia'
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = DARK_RED
    return h

def add_h2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = 'Georgia'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE
    return h

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r_bold = p.add_run(bold_prefix)
    r_bold.font.bold = True
    p.add_run(text)
    return p

out_dir = "c:/Users/dorel/OneDrive/Desktop/RogueLike_game/docs"
os.makedirs(out_dir, exist_ok=True)

# DOC 6: UI & UX DESIGN
doc6 = create_base_doc("KAGE NO YOKAI: 1200 AD\nГлава VI: Дизайн Пользовательского Интерфейса (UI / UX)", "Дизайн-Документ Экранов Главного Меню, Настроек, Выбора Скинов, Выбора Режимов и Карточек Левел-апа")

add_h1(doc6, "1. Концепция Пользовательского Интерфейса (UI / UX)")
p = doc6.add_paragraph()
p.add_run("Интерфейс игры выполнен в едином японо-самурайском пиксельном стиле. Использованы текстуры соснового дерева, свитков рисовой бумаги, позолоченные рамки и эффекты частиц (падающие лепестки сакуры, пламя фонарей).")

add_h2(doc6, "Разбор 5 Ключевых Экранов Игры:")

ui_screens = [
    ("1. Экран Главного Меню (Main Menu)", "При запуске открывается 3D Pixel задний план пагоды с фонарями. Кнопки: «В БОЙ», «ПЕРСОНАЖИ И СКИНЫ», «ВЫБОР ЭТАПА И РЕЖИМА», «НАСТРОЙКИ», «ВЫХОД»."),
    ("2. Раздел Настроек (Settings)", "Оверлей на фоне деревянного свитка. Вкладки: Звук (громкость/эффекты), Графика (шейдер пикселизации, bloom), Управление (WASD/мышь), Язык."),
    ("3. Выбор Скинов и Персонажей (Skin Selection)", "В центре 3D-модель воина в самурайских доспехах. Справа панель характеристик (HP, урон, пассивка). Внизу карусель открытых и заблокированных скинов. Вверху полоса XP аккаунта."),
    ("4. Выбор Этапа и Режима (Campaign & Endless)", "Переключатель режимов: «12 ЭТАПОВ КАМПАНИИ» и «БЕСКОНЕЧНЫЙ РЕЖИМ (Endless)». Сетка карт с миниатюрами, список врагов, аватара босса и кнопка «НАЧАТЬ ЗАБЕГ»."),
    ("5. Карточки Повышения Уровня (Level-Up Upgrade Screen)", "При паузе во время игры выводится 3 драфт-карточки со свечением редкости (Обычное, Редкое, Легендарное), кнопка «РЕРОЛЛ (1)» и «ПРОПУСК». Вверху индикаторы слотов 3/3 оружия и 3/3 способностей.")
]

for name, desc in ui_screens:
    add_h2(doc6, name)
    p_desc = doc6.add_paragraph()
    p_desc.add_run(desc)

doc6.save(f"{out_dir}/GDD_6_UI_UX_Design.docx")
print("GDD Document 6 UI/UX generated successfully!")
